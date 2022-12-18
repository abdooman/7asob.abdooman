import docx
from pathlib import Path

doc = docx.Document(Path.home() / Path('OneDrive', 'Word', 'academy_1.docx'))
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)

print(len(doc.paragraphs[2].runs))
print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)

print('#' * 50)

def  getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getText(Path.home() / Path('OneDrive', 'Word', 'academy_1.docx')))
    