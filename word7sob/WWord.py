import docx 
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")

# إضافة نص باللغة العربية 
paragraph = mydoc.add_paragraph("هذه هي الفقرة الثالثة ")
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # تحويل كتابة الملف من جهة اليسار الى اليمين
mydoc.save(Path.home() / Path('OneDrive', 'WR.docx'))

third_paragraph = mydoc.add_paragraph("this is the third paragraph.")
third_paragraph.add_run(" this is a section at the end of third paragraph")


#add header اضافة عنوان 
mydoc.add_heading("THis is level 1 heading", 0)
mydoc.add_heading("THis is level 2 heading", 1)
mydoc.add_heading("THis is level 3 heading", 2)
mydoc.add_heading("THis is level 4 heading", 3)
mydoc.add_heading("THis is level 5 heading", 4)


#style
print(mydoc.paragraphs[0].style)
print(mydoc.paragraphs[3].style)
print(mydoc.paragraphs[5].style)
mydoc.paragraphs[0].style = mydoc.styles['Heading 1']
print(mydoc.paragraphs[0].style)

mydoc.add_paragraph("This is style paragraph.", "Title")

#image
mydoc.add_picture(str(Path.home() / Path('OneDrive', 'Screenshot_2.png')), width=docx.shared.Inches(5), height=docx.shared.Inches(7))  # type: ignore

mydoc.save(Path.home() / Path('OneDrive', 'WR.docx')) # يجب حفظ الملف بعد التعديل عليه
